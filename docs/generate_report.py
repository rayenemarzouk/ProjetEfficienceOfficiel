from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
path = r'c:\ProjetEfficienceOfficiel\docs\rapport_pfe_marzouk_rayene_final.docx'
doc = Document()

title = doc.add_paragraph()
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
title_run = title.add_run('Rapport PFE - Projet Efficience Officiel')
title_run.bold = True
title_run.font.size = Pt(24)

doc.add_paragraph()
doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
sub_run = sub.add_run('Mise à jour technique, intégration des captures et déploiement')
sub_run.font.size = Pt(14)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
info.add_run('Date : Juin 2026')

doc.add_page_break()

doc.add_heading('Résumé', level=1)
doc.add_paragraph('Ce rapport présente l’état actuel du projet Efficience Officiel, la mise à jour des composants techniques, l’intégration de la fonctionnalité d’import de données, la gestion des rôles ainsi que la stratégie de déploiement pour un hébergement frontend statique sur Hostinger et un backend Node.js séparé.')

doc.add_heading('1. Introduction', level=1)
doc.add_paragraph('1.1 Contexte')
doc.add_paragraph('Le projet Efficience Officiel est une solution de suivi analytique pour cabinets et consultants. Il combine un dashboard React, une API Node.js/Express, et une base de données MongoDB Atlas.')

doc.add_paragraph('1.2 Objectifs')
doc.add_paragraph('L’objectif de cette mise à jour est de documenter précisément le projet, d’ajouter la fonctionnalité d’import de fichiers, de formaliser la gestion des rôles, et de préparer le déploiement sur une infrastructure moderne avec Hostinger pour le frontend et un backend Node.js hébergé séparément.')

doc.add_paragraph('1.3 Portée')
doc.add_paragraph('Le rapport couvre l’architecture globale, les composants frontend et backend, les flux d’import de données, les contraintes de sécurité, les tâches planifiées (cron), et le processus de déploiement.')

doc.add_heading('2. Architecture du système', level=1)
doc.add_paragraph('2.1 Architecture générale')
doc.add_paragraph('L’application est structurée en trois couches : le frontend React/Vite/Tailwind, le backend Express/Mongoose, et la base de données MongoDB Atlas. La communication entre le frontend et le backend se fait via une API REST sécurisée par JWT.')

doc.add_paragraph('2.2 Frontend')
doc.add_paragraph('Le frontend est développé avec React et Vite. Le client Axios est centralisé dans frontend/src/services/api.js, avec une variable dynamique VITE_API_URL || "/api" pour assurer la flexibilité de déploiement.')

doc.add_paragraph('2.3 Backend')
doc.add_paragraph('Le backend utilise Node.js et Express. Le serveur principal est défini dans backend/server.js. Il propose des routes pour l’authentification, les imports de données, les rapports, et la gestion des utilisateurs.')

doc.add_paragraph('2.4 Base de données')
doc.add_paragraph('La persistance est assurée par MongoDB Atlas via Mongoose. Les modèles principaux se trouvent dans backend/models/* et couvrent les analyses, les utilisateurs, les rapports, et les paramètres de l’application.')

doc.add_heading('3. Fonctionnalités principales', level=1)
doc.add_paragraph('3.1 Authentification et rôles')
doc.add_paragraph('L’authentification est gérée par email/mot de passe, avec bcrypt et JWT. Le middleware backend/middleware/auth.js protège les routes et applique le contrôle de rôle : admin, consultant, practitioner.')

doc.add_paragraph('3.2 Gestion des données et import')
doc.add_paragraph('La page de gestion des données permet d’importer des fichiers de différents types : réalisation, rendez-vous, jours ouverts, devis, et en-cours. La fonctionnalité est implémentée par frontend/src/pages/practitioner/DataManagement.jsx et le backend backend/routes/data.js.')

doc.add_paragraph('3.3 Génération de rapports')
doc.add_paragraph('Le projet inclut une génération de rapports PDF via backend/services/pdfGenerator.js. La génération est désactivée en environnement cloud si Puppeteer n’est pas activé.')

doc.add_paragraph('3.4 Tâches planifiées')
doc.add_paragraph('Le backend initialise une tâche planifiée à l’aide de initCronJobs() dans backend/server.js. Cette fonction se trouve dans backend/services/cronJobs.js et orchestre les traitements automatisés des indicateurs.')

doc.add_heading('4. Implémentation détaillée', level=1)
doc.add_paragraph('4.1 Frontend')
doc.add_paragraph('Le frontend expose plusieurs pages : admin, consultant, practitioner. Les composants de visualisation et les dashboards sont organisés dans frontend/src/components et frontend/src/pages.')

doc.add_paragraph('4.2 Backend')
doc.add_paragraph('Le backend sert l’API et, en production, les fichiers statiques du frontend présents dans backend/public. Il configure également CORS pour autoriser les origines locales et l’URL Hostinger actuelle.')

doc.add_paragraph('4.3 Imports de données')
doc.add_paragraph('Le point d’entrée backend pour l’import est POST /api/data/import/:type. La requête attend un fichier multipart/form-data avec le champ file et un header Authorization: Bearer <token>. En frontend, la sélection de fichier est déclenchée par le composant DataManagement.jsx.')

doc.add_paragraph('4.4 Sécurité')
doc.add_paragraph('Le serveur vérifie le JWT sur chaque route protégée. Le code prend en charge la restriction de certaines opérations aux administrateurs uniquement.')

doc.add_heading('5. Captures d’écran et légendes', level=1)
doc.add_paragraph('Les captures fournies ont été intégrées dans le rapport sous forme de placeholders. Il conviendra d’insérer les images finales au moment de la mise en page.')
for text in ['Figure 1 : Interface de sélection et d’import', 'Figure 2 : Message de retour après import', 'Figure 3 : Requête POST /api/data/import/{type} et réponse JSON', 'Figure 4 : Écran de connexion (Login)', 'Figure 5 : Dashboard analytique principal', 'Figure 6 : Gestion des cabinets et actions', 'Figure 7 : Modèles IA intégrés', 'Figure 8 : Génération de rapport PDF', 'Figure 9 : Initialisation Cron dans server.js', 'Figure 10 : Déploiement Hostinger']:
    doc.add_paragraph(text)
doc.add_heading('6. Déploiement', level=1)
doc.add_paragraph('6.1 Stratégie de déploiement')
doc.add_paragraph('La stratégie retenue sépare le frontend statique du backend API. Le frontend est déployé sur Hostinger en tant que site statique ou contenu uploadé. Le backend Node.js est hébergé sur un service cloud compatible et communique avec le frontend via une API HTTPS.')
doc.add_paragraph('6.2 Variables d’environnement')
doc.add_paragraph('Les variables recommandées sont : MONGODB_URI, JWT_SECRET, EMAIL_USER, EMAIL_PASS, PORT, NODE_ENV=production, PUPPETEER_ENABLED=false, et VITE_API_URL pointant vers l’URL du backend.')
doc.add_paragraph('6.3 CORS')
doc.add_paragraph('La configuration CORS actuelle permet les origines locales et l’URL Hostinger. Elle garantit qu’un site statique hébergé sur Hostinger peut appeler l’API Backend.')
doc.add_heading('7. Conclusion', level=1)
doc.add_paragraph('Ce rapport finalise la mise à jour du projet Efficience Officiel en intégrant la documentation technique, l’architecture, les fonctionnalités métier, la gestion de l’import, et la préparation du déploiement. Le produit est prêt pour un hébergement production avec un frontend statique et un backend API sécurisé.')
doc.save(path)
print(path)
